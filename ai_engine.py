import re
import random
import pdfplumber
import database as db

try:
    import docx as _docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import nltk
    from nltk.tokenize import word_tokenize
    from nltk.stem import WordNetLemmatizer

    for path, pkg in [("tokenizers/punkt_tab", "punkt_tab"), ("corpora/wordnet", "wordnet"),
                       ("corpora/omw-1.4", "omw-1.4")]:
        try:
            nltk.data.find(path)
        except LookupError:
            nltk.download(pkg, quiet=True)

    _lemmatizer = WordNetLemmatizer()
    NLTK_AVAILABLE = True
except Exception:
    NLTK_AVAILABLE = False

try:
    from sentence_transformers import SentenceTransformer, util
    _model = SentenceTransformer("all-MiniLM-L6-v2")
    BERT_AVAILABLE = True
except Exception:
    BERT_AVAILABLE = False


# =========================== Resume Parsing ===========================

def extract_text_pdf(pdf_file):
    text = ""
    with pdfplumber.open(pdf_file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    return text


def extract_text_docx(docx_file):
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx isn't installed. Run: pip install python-docx")

    document = _docx.Document(docx_file)
    parts = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    parts.append(cell.text)
    return "\n".join(parts)


def extract_text(resume_file, filename=None):
    name = filename or getattr(resume_file, "name", "") or ""
    if name.lower().endswith(".docx"):
        return extract_text_docx(resume_file)
    return extract_text_pdf(resume_file)


NAME_SKIP_WORDS = ["curriculum vitae", "resume", "bio data", "biodata", "contact",
                    "profile", "summary", "objective", "personal details", "address"]


def extract_name(text):
    for line in text.split("\n"):
        line = line.strip()
        if not line:
            continue
        lowered = line.lower()
        if any(w in lowered for w in NAME_SKIP_WORDS):
            continue
        letters = sum(c.isalpha() for c in line)
        if letters < len(line) * 0.6:
            continue
        if len(line.split()) >= 2 and len(line) < 35:
            return line
    return "Unknown"


def extract_email(text):
    m = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    return m.group() if m else "Not Found"


def extract_phone(text):
    m = re.search(r"(\+91[- ]?)?[6-9]\d{9}", text)
    return m.group() if m else "Not Found"


# Each degree maps to a regex pattern (word-boundary safe, dots/spaces optional)
# and a level used for scoring. Patterns are checked with re.search, never with
# plain substring "in" checks, to avoid false positives like "be" matching
# inside "based"/"best"/"become".
DEGREE_PATTERNS = [
    ("phd",       r"\bph\.?\s?d\.?\b", 4),
    ("doctorate", r"\bdoctorate\b", 4),
    ("master",    r"\bmaster'?s?\b", 3),
    ("m.tech",    r"\bm\.?\s?tech\b", 3),
    ("mba",       r"\bm\.?\s?b\.?\s?a\.?\b", 3),
    ("m.sc",      r"\bm\.?\s?sc\.?\b", 3),
    ("mca",       r"\bm\.?\s?c\.?\s?a\.?\b", 3),
    ("m.com",     r"\bm\.?\s?com\.?\b", 3),
    ("bachelor",  r"\bbachelor'?s?\b", 2),
    ("b.tech",    r"\bb\.?\s?tech\b", 2),
    ("b.e",       r"\bb\.?\s?e\.?\b(?!\s?d)", 2),
    ("b.sc",      r"\bb\.?\s?sc\.?\b", 2),
    ("bca",       r"\bb\.?\s?c\.?\s?a\.?\b", 2),
    ("b.com",     r"\bb\.?\s?com\.?\b", 2),
    ("diploma",   r"\bdiploma\b", 1),
]


def extract_education(text):
    text = text.lower()
    found = [name for name, pattern, _ in DEGREE_PATTERNS if re.search(pattern, text)]
    return ", ".join(found) if found else "Not Mentioned"


def parse_resume(resume_file):
    text = extract_text(resume_file)
    return {
        "name": extract_name(text), "email": extract_email(text),
        "phone": extract_phone(text), "education": extract_education(text),
        "resume_text": text,
    }


# =========================== Skill Extraction ===========================

SKILL_CATEGORIES = {
    "Programming": ["python", "java", "c", "c++", "c#", "javascript", "typescript",
                     "go", "rust", "r", "php", "ruby", "kotlin", "swift", "scala"],
    "Machine Learning & AI": ["machine learning", "deep learning", "artificial intelligence",
                               "neural networks", "computer vision", "reinforcement learning",
                               "tensorflow", "pytorch", "keras", "scikit-learn", "xgboost",
                               "opencv", "huggingface", "transformers", "generative ai"],
    "Cloud": ["aws", "azure", "gcp", "google cloud platform", "cloud computing", "ec2", "s3",
              "lambda", "cloudformation", "azure devops", "google cloud functions", "cloud architecture"],
    "Databases": ["sql", "mysql", "postgresql", "mongodb", "oracle", "sqlite", "redis",
                   "cassandra", "dynamodb", "nosql", "database design", "data warehousing"],
    "NLP": ["nlp", "sentiment analysis", "text mining", "named entity recognition", "tokenization",
             "bert", "spacy", "nltk", "language modeling", "topic modeling"],
    "DevOps": ["docker", "kubernetes", "jenkins", "ci/cd", "git", "github", "gitlab",
                "ansible", "terraform", "linux", "bash scripting", "devops"],
    "Analytics": ["data analysis", "data visualization", "tableau", "power bi", "excel",
                   "statistics", "pandas", "numpy", "matplotlib", "seaborn",
                   "business intelligence", "data mining"],
    "Web": ["html", "css", "react", "angular", "vue", "nodejs", "express", "django",
             "flask", "rest api", "graphql", "bootstrap", "spring boot", "spring",
             "microservices"],
    "Soft Skills": ["communication", "leadership", "teamwork", "problem solving",
                      "time management", "critical thinking", "adaptability",
                      "collaboration", "creativity", "project management"],
}

SKILL_TO_CATEGORY = {s: c for c, skills in SKILL_CATEGORIES.items() for s in skills}
MASTER_SKILLS = sorted(SKILL_TO_CATEGORY.keys())

SYNONYMS = {
    "ml": "machine learning", "ai": "artificial intelligence", "dl": "deep learning",
    "rl": "reinforcement learning", "genai": "generative ai",
    "gen ai": "generative ai", "llm": "transformers", "hugging face": "huggingface",
    "k8s": "kubernetes", "cicd": "ci/cd", "continuous integration": "ci/cd",
    "continuous deployment": "ci/cd", "js": "javascript", "ts": "typescript", "py": "python",
    "golang": "go", "postgres": "postgresql", "mongo": "mongodb", "tf": "tensorflow",
    "sklearn": "scikit-learn", "scikit learn": "scikit-learn", "powerbi": "power bi",
    "rest apis": "rest api", "restful api": "rest api", "node.js": "nodejs", "node": "nodejs",
    "reactjs": "react", "react.js": "react", "vuejs": "vue", "vue.js": "vue",
    "gcp cloud": "gcp", "google cloud": "google cloud platform",
    "db design": "database design", "data viz": "data visualization",
}


# Common English phrases that would otherwise trigger false-positive matches
# for short/ambiguous skill tokens ("r", "c", "go") once punctuation is
# stripped by normalize_text (e.g. "R&D" -> "r d" -> looks like the token "r").
# These are scrubbed out of a copy of the raw text before skill matching runs.
_AMBIGUOUS_FALSE_POSITIVE_PHRASES = [
    r"r\s*&\s*d\b", r"r\s+and\s+d\b", r"research\s+and\s+development",
    r"go\s+the\s+extra\s+mile", r"go\s+above\s+and\s+beyond", r"go\s+beyond",
    r"willing\s+to\s+go\b", r"go\s+through\b", r"go\s+over\b",
    r"class\s+c\b", r"vitamin\s+c\b", r"grade\s+c\b", r"plan\s+c\b",
]


def _scrub_ambiguous_phrases(text):
    text = text.lower()
    for pattern in _AMBIGUOUS_FALSE_POSITIVE_PHRASES:
        text = re.sub(pattern, " ", text)
    return text


def normalize_text(text):
    text = text.lower()
    if NLTK_AVAILABLE:
        try:
            text = " ".join(_lemmatizer.lemmatize(t) for t in word_tokenize(text))
        except Exception:
            pass
    text = re.sub(r"[^a-z0-9+#./ ]", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def _contains_token(text, token):
    if token == "c++":
        pattern = r"c\+\+"
    elif token == "c#":
        pattern = r"c#"
    elif "/" in token or "." in token:
        pattern = re.escape(token)
    else:
        pattern = r"\b" + re.escape(token) + r"\b"
    return re.search(pattern, text) is not None


def extract_skills(text):
    normalized = normalize_text(_scrub_ambiguous_phrases(text))
    found = {s for s in MASTER_SKILLS if _contains_token(normalized, s)}
    found |= {canon for variant, canon in SYNONYMS.items() if _contains_token(normalized, variant)}
    return sorted(found)


def extract_skills_by_category(text):
    flat = extract_skills(text)
    grouped = {c: [] for c in SKILL_CATEGORIES}
    for skill in flat:
        grouped[SKILL_TO_CATEGORY[skill]].append(skill)
    return grouped


def analyze_skills(candidate, job_description):
    resume_skills = extract_skills(candidate["resume_text"])
    job_skills = extract_skills(job_description)

    matched = [s for s in job_skills if s in resume_skills]
    missing = [s for s in job_skills if s not in resume_skills]
    percentage = 100 if not job_skills else round(len(matched) / len(job_skills) * 100, 2)

    recommendations = (["Excellent! Your resume covers all required skills."] if not missing else
                        [f"Consider learning '{s}' to improve your ATS score." for s in missing])

    return {
        "Resume Skills": resume_skills, "Job Skills": job_skills,
        "Matched Skills": matched, "Missing Skills": missing,
        "Skill Match Percentage": percentage, "Recommendations": recommendations,
        "Resume Skills By Category": extract_skills_by_category(candidate["resume_text"]),
        "Job Skills By Category": extract_skills_by_category(job_description),
    }


# =========================== BERT Scoring ===========================

def _word_overlap(a, b):
    words = lambda t: set(re.findall(r"[a-z0-9]+", t.lower()))
    aw, bw = words(a), words(b)
    return round(min(len(aw & bw) / len(bw), 1.0) * 100, 2) if bw else 0.0


def _rescale_similarity(raw_cos_sim):
    # Raw cosine similarity from all-MiniLM-L6-v2 between a resume and a job
    # description rarely approaches 1.0, even for a genuinely excellent match,
    # because the two texts differ in length/structure/style. Empirically,
    # ~0.25 cosine similarity is roughly "unrelated" and ~0.75+ is roughly
    # "excellent match" for this model on resume-vs-JD text pairs. Treating
    # the raw score linearly as a 0-100 percentage compresses every score
    # into a low ceiling, so we rescale against these realistic bounds instead.
    floor, ceiling = 0.25, 0.75
    scaled = (raw_cos_sim - floor) / (ceiling - floor)
    return round(max(0.0, min(scaled, 1.0)) * 100, 2)


def semantic_similarity(resume_text, job_description):
    if not BERT_AVAILABLE:
        return _word_overlap(resume_text, job_description)
    r = _model.encode(resume_text, convert_to_tensor=True)
    j = _model.encode(job_description, convert_to_tensor=True)
    return _rescale_similarity(util.cos_sim(r, j).item())


def bulk_semantic_similarity(resume_texts, job_description, use_cache=True):
    if not resume_texts:
        return []
    if not BERT_AVAILABLE:
        return [_word_overlap(t, job_description) for t in resume_texts]

    all_texts = resume_texts + [job_description]
    cached = db.get_cached_embeddings(all_texts) if use_cache else {}
    to_encode = [t for t in all_texts if t not in cached]

    if to_encode:
        fresh = _model.encode(to_encode, convert_to_tensor=True, batch_size=32, show_progress_bar=False)
        new_cache = {t: fresh[i].tolist() for i, t in enumerate(to_encode)}
        if use_cache:
            db.store_embeddings(new_cache)
        cached.update(new_cache)

    import torch
    jd_emb = torch.tensor(cached[job_description])
    resume_embs = torch.tensor([cached[t] for t in resume_texts])
    sims = util.cos_sim(resume_embs, jd_emb).squeeze(-1)
    if sims.dim() == 0:
        sims = sims.unsqueeze(0)
    return [_rescale_similarity(s.item()) for s in sims]


# =========================== Experience & Education ===========================

YEARS_PATTERN = re.compile(r"(\d+(?:\.\d+)?)\+?\s*(?:years|yrs|year)\b(?:\s*of\s*experience)?", re.IGNORECASE)
EXPERIENCE_KEYWORDS = ["internship", "intern", "experience", "project", "developer", "engineer", "research", "freelance"]
# Reuses the same word-boundary-safe DEGREE_PATTERNS defined near extract_education,
# instead of the old plain-substring dict (which had "be": 2 and matched inside
# "based"/"best"/"become", silently treating almost any text as requiring a degree).


def extract_years_of_experience(text):
    matches = YEARS_PATTERN.findall(text)
    return max(float(m) for m in matches) if matches else None


def experience_score(candidate, job_description):
    resume_text = candidate["resume_text"]
    candidate_years = extract_years_of_experience(resume_text)
    required_years = extract_years_of_experience(job_description)

    if candidate_years is not None and required_years:
        return round(min(candidate_years / required_years, 1.0) * 100, 2)
    if candidate_years is not None:
        return round(min(candidate_years / 10, 1.0) * 100, 2)

    text = resume_text.lower()
    count = sum(1 for w in EXPERIENCE_KEYWORDS if w in text)
    return min(count * 10, 60)


def _degree_level(text):
    text = text.lower()
    return max((lvl for _, pattern, lvl in DEGREE_PATTERNS if re.search(pattern, text)), default=0)


def education_score(candidate, job_description):
    candidate_level = _degree_level(candidate["education"])
    required_level = _degree_level(job_description)

    if required_level > 0:
        if candidate_level >= required_level:
            return 100
        return 60 if candidate_level == required_level - 1 else 30

    return {4: 100, 3: 90, 2: 80, 1: 60, 0: 40}[candidate_level]


# =========================== Verdict & Final Score ===========================

DEFAULT_WEIGHTS = {"semantic": 40, "skill": 35, "experience": 15, "education": 10}


def generate_verdict(score):
    if score >= 85:
        return "Strong Fit — Highly Recommended for Interview"
    if score >= 70:
        return "Good Fit — Recommended for Interview"
    if score >= 50:
        return "Moderate Fit — Consider with Reservations"
    return "Weak Fit — Not Recommended at This Time"


def generate_ai_feedback(candidate, ats_result, skill_report):
    lines = [generate_verdict(ats_result["Final Score"]), ""]
    matched, missing = skill_report["Matched Skills"], skill_report["Missing Skills"]

    if matched:
        lines.append(f"Strengths: {candidate['name']} matches {len(matched)} of the skills the role "
                      f"asks for, including {', '.join(matched[:5])}.")
    if missing:
        lines.append(f"Skill Gap: missing {len(missing)} requested skill(s) — {', '.join(missing)}. "
                      f"Consider probing these in interview or flagging as upskilling areas.")
    else:
        lines.append("No skill gaps detected against this job description.")

    if ats_result["Experience Score"] < 60:
        lines.append("Experience appears below what the role typically expects — verify project depth and duration during screening.")
    if ats_result["Education Score"] < 60:
        lines.append("Education level is below the role's stated requirement.")

    return "\n".join(lines)


def _weighted_final(semantic, skill, experience, education, w):
    return round(semantic * w["semantic"] / 100 + skill * w["skill"] / 100 +
                 experience * w["experience"] / 100 + education * w["education"] / 100, 2)


def calculate_ats_score(candidate, job_description, skill_report, weights=None):
    w = weights or DEFAULT_WEIGHTS
    semantic = semantic_similarity(candidate["resume_text"], job_description)
    skill = skill_report["Skill Match Percentage"]
    education = education_score(candidate, job_description)
    experience = experience_score(candidate, job_description)

    result = {"Final Score": _weighted_final(semantic, skill, experience, education, w),
              "Semantic Score": semantic, "Skill Score": skill,
              "Education Score": education, "Experience Score": experience}
    result["Verdict"] = generate_verdict(result["Final Score"])
    result["AI Feedback"] = generate_ai_feedback(candidate, result, skill_report)
    return result


def bulk_calculate_ats_scores(candidates, job_description, skill_reports, weights=None):
    w = weights or DEFAULT_WEIGHTS
    semantic_scores = bulk_semantic_similarity([c["resume_text"] for c in candidates], job_description)

    results = []
    for candidate, skill_report, semantic in zip(candidates, skill_reports, semantic_scores):
        skill = skill_report["Skill Match Percentage"]
        education = education_score(candidate, job_description)
        experience = experience_score(candidate, job_description)

        result = {"Final Score": _weighted_final(semantic, skill, experience, education, w),
                   "Semantic Score": semantic, "Skill Score": skill,
                   "Education Score": education, "Experience Score": experience}
        result["Verdict"] = generate_verdict(result["Final Score"])
        result["AI Feedback"] = generate_ai_feedback(candidate, result, skill_report)
        results.append(result)

    return results


# =========================== Interview Questions ===========================

TECHNICAL_QUESTIONS = {
    "python": [("Easy", "What is the difference between a list and a tuple in Python?"),
                ("Medium", "Explain Python's GIL and how it affects multithreading."),
                ("Hard", "How would you optimize a Python function that's a performance bottleneck?")],
    "sql": [("Easy", "What is the difference between INNER JOIN and LEFT JOIN?"),
             ("Medium", "How would you find duplicate rows in a table using SQL?"),
             ("Hard", "Explain how you would optimize a slow-running SQL query with a large dataset.")],
    "machine learning": [("Easy", "What is the difference between supervised and unsupervised learning?"),
                           ("Medium", "How do you handle overfitting in a machine learning model?"),
                           ("Hard", "Walk through how you would design an end-to-end ML pipeline for a new problem.")],
    "deep learning": [("Medium", "What is the vanishing gradient problem and how is it addressed?"),
                        ("Hard", "Compare CNNs and Transformers for a vision task — when would you pick one over the other?")],
    "nlp": [("Easy", "What is tokenization and why is it needed in NLP?"),
             ("Medium", "Explain how BERT differs from earlier word-embedding approaches like Word2Vec.")],
    "docker": [("Easy", "What is the difference between a Docker image and a container?"),
                ("Medium", "How would you reduce the size of a Docker image?")],
    "git": [("Easy", "What is the difference between `git merge` and `git rebase`?")],
    "react": [("Easy", "What is the difference between props and state in React?"),
               ("Medium", "Explain how the virtual DOM improves rendering performance.")],
    "aws": [("Medium", "How would you design a highly available architecture on AWS?")],
}

GENERIC_TECHNICAL = [
    ("Medium", "Describe a project where you used {skill}. What challenges did you face?"),
    ("Easy", "How would you explain {skill} to someone with no technical background?"),
]

HR_QUESTIONS = [
    ("Easy", "Tell me about yourself and why you're interested in this role."),
    ("Easy", "What are your greatest strengths and weaknesses?"),
    ("Medium", "Describe a time you disagreed with a teammate. How did you resolve it?"),
    ("Medium", "Where do you see yourself in the next 3-5 years?"),
    ("Hard", "Tell me about a time you failed at something important. What did you learn?"),
]


def generate_interview_questions(matched_skills, num_technical=5, num_hr=3):
    questions = []
    for skill in matched_skills:
        if skill in TECHNICAL_QUESTIONS:
            for difficulty, question in TECHNICAL_QUESTIONS[skill]:
                questions.append({"type": f"Technical ({skill.title()})", "difficulty": difficulty, "question": question})
        else:
            difficulty, template = random.choice(GENERIC_TECHNICAL)
            questions.append({"type": f"Technical ({skill.title()})", "difficulty": difficulty,
                               "question": template.format(skill=skill)})

    seen, deduped = set(), []
    for q in questions:
        if q["question"] not in seen:
            seen.add(q["question"])
            deduped.append(q)

    technical = deduped[:num_technical] or [{"type": "Technical (General)", "difficulty": "Easy",
                                              "question": "Walk me through a recent technical project you're proud of."}]
    hr = [{"type": "HR", "difficulty": d, "question": q} for d, q in HR_QUESTIONS[:num_hr]]
    return technical + hr
